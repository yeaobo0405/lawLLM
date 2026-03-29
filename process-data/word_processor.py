"""
Word文档预处理模块
支持 .doc 和 .docx 格式
"""
import os
import re
import logging
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, asdict

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """预处理后的文档结构"""
    file_name: str
    file_path: str
    doc_type: str
    law_name: str = ""
    chapter: str = ""
    article_number: str = ""
    case_number: str = ""
    judgment_date: str = ""
    case_type: str = ""
    content: str = ""
    articles: List[Dict] = None
    
    def __post_init__(self):
        if self.articles is None:
            self.articles = []


class WordProcessor:
    """Word文档处理器"""
    
    def __init__(self):
        self.header_patterns = [
            r'^第\s*\d+\s*页',
            r'^\d+\s*/\s*\d+',
            r'^\s*\d+\s*$',
        ]
        self.footer_patterns = [
            r'^\s*第\s*\d+\s*页\s*$',
            r'^\s*\d+\s*/\s*\d+\s*$',
            r'^\s*页\s*$',
        ]
        
    def detect_file_format(self, file_path: str) -> str:
        """检测文件实际格式"""
        with open(file_path, 'rb') as f:
            header = f.read(4)
        
        if header == b'\x50\x4b\x03\x04':
            return 'docx'
        elif header == b'\xd0\xcf\x11\xe0':
            return 'doc'
        else:
            return 'unknown'
    
    def extract_text_docx(self, file_path: str) -> str:
        """从DOCX提取纯文本（剔除页眉页脚批注等）"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            text_parts = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and not self._is_header_footer(text):
                    text_parts.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"python-docx提取失败: {e}")
            return ""
    
    def extract_text_doc(self, file_path: str) -> str:
        """从DOC提取纯文本（使用COM接口）"""
        try:
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(os.path.abspath(file_path))
            
            text_parts = []
            
            for para in doc.Paragraphs:
                text = para.Range.Text.strip()
                if text and not self._is_header_footer(text):
                    text_parts.append(text)
            
            for table in doc.Tables:
                for row in table.Rows:
                    row_text = []
                    for cell in row.Cells:
                        cell_text = cell.Range.Text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            doc.Close(False)
            word.Quit()
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"COM方式提取DOC失败: {e}")
            return ""
    
    def _is_header_footer(self, text: str) -> bool:
        """判断是否为页眉页脚内容"""
        text = text.strip()
        
        for pattern in self.header_patterns + self.footer_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        if len(text) < 5 and text.isdigit():
            return True
        
        return False
    
    def process(self, file_path: str, doc_type: str = "law") -> ProcessedDocument:
        """处理Word文档"""
        file_name = os.path.basename(file_path)
        
        file_format = self.detect_file_format(file_path)
        
        if file_format == 'docx':
            text = self.extract_text_docx(file_path)
        elif file_format == 'doc':
            text = self.extract_text_doc(file_path)
        else:
            logger.warning(f"未知文件格式: {file_path}")
            text = ""
        
        if not text:
            return ProcessedDocument(
                file_name=file_name,
                file_path=file_path,
                doc_type=doc_type
            )
        
        processed_doc = ProcessedDocument(
            file_name=file_name,
            file_path=file_path,
            doc_type=doc_type
        )
        
        if doc_type == "law":
            self._extract_law_metadata(text, processed_doc)
        else:
            self._extract_case_metadata(text, processed_doc)
        
        processed_doc.content = text
        
        return processed_doc
    
    def _extract_law_metadata(self, text: str, doc: ProcessedDocument):
        """提取法律文档元数据"""
        file_base = os.path.splitext(doc.file_name)[0]
        if '_' in file_base:
            doc.law_name = file_base.split('_')[0]
        else:
            law_name_match = re.search(r'《(.+?)》', text[:1000])
            if law_name_match:
                doc.law_name = law_name_match.group(1)
            else:
                doc.law_name = file_base
        
        chapter_match = re.search(r'第([一二三四五六七八九十]+)[章节]', text)
        if chapter_match:
            doc.chapter = chapter_match.group(0)
        
        article_pattern = r'第([一二三四五六七八九十百零]+)条'
        matches = list(re.finditer(article_pattern, text))
        
        for i, match in enumerate(matches):
            article_num = self._chinese_to_number(match.group(1))
            start = match.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            article_text = text[start:end].strip()
            
            doc.articles.append({
                "article_number": article_num,
                "content": article_text
            })
    
    def _extract_case_metadata(self, text: str, doc: ProcessedDocument):
        """提取案例文档元数据"""
        case_number_match = re.search(r'[（(](\d{4})[)）].*?第?(\d+号?)', text)
        if case_number_match:
            doc.case_number = case_number_match.group(0)
        
        date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', text)
        if date_match:
            doc.judgment_date = date_match.group(0)
        
        case_types = ['民事', '刑事', '行政', '经济', '知识产权', '劳动争议']
        for ct in case_types:
            if ct in text:
                doc.case_type = ct
                break
    
    def _chinese_to_number(self, chinese: str) -> str:
        """中文数字转阿拉伯数字"""
        mapping = {
            '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
            '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
            '十': 10, '百': 100
        }
        
        if len(chinese) == 1:
            return str(mapping.get(chinese, chinese))
        
        result = 0
        temp = 0
        
        for char in chinese:
            if char in mapping:
                val = mapping[char]
                if val >= 10:
                    if temp == 0:
                        temp = 1
                    result += temp * val
                    temp = 0
                else:
                    temp = val
            else:
                return chinese
        
        result += temp
        return str(result) if result > 0 else chinese
