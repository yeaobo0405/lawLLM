"""
文档处理模块
使用LangChain实现PDF、TXT、DOCX格式文档的加载、分段、清洗、结构化处理
"""
import os
import re
import logging
import cn2an
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass

from langchain_community.document_loaders import Docx2txtLoader, PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

try:
    from ..config import settings
except ImportError:
    from config import settings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class LegalMetadata:
    """法条元数据结构"""
    law_name: str
    chapter: str
    article_number: str
    file_path: str
    file_name: str
    doc_type: str = "law"


@dataclass
class CaseMetadata:
    """案例元数据结构"""
    case_number: str
    judgment_date: str
    case_type: str
    file_path: str
    file_name: str
    doc_type: str = "case"


class DocumentProcessor:
    """
    文档处理器
    负责加载、分段、清洗、结构化处理法律文档
    """
    
    def __init__(self, knowledge_base_path: str = "./knowledge_base"):
        """
        初始化文档处理器
        
        Args:
            knowledge_base_path: 知识库文档存储路径
        """
        self.knowledge_base_path = knowledge_base_path
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", "；", "，", " "]
        )
        
    def load_docx(self, file_path: str) -> List[Document]:
        """
        加载DOCX格式文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            Document对象列表
        """
        try:
            loader = Docx2txtLoader(file_path)
            documents = loader.load()
            
            if not documents or not documents[0].page_content.strip():
                logger.warning(f"DOCX文档内容为空: {file_path}")
                return self._load_docx_fallback(file_path)
            
            logger.info(f"成功加载DOCX文档: {file_path}, 页数: {len(documents)}")
            return documents
        except Exception as e:
            logger.warning(f"标准方式加载DOCX失败: {file_path}, 尝试备用方法...")
            return self._load_docx_fallback(file_path)
    
    def _load_docx_fallback(self, file_path: str) -> List[Document]:
        """
        备用方法加载DOCX文档（处理损坏或伪装的文件）
        
        Args:
            file_path: 文档路径
            
        Returns:
            Document对象列表
        """
        with open(file_path, 'rb') as f:
            header = f.read(4)
        
        is_ole = header == b'\xd0\xcf\x11\xe0'
        
        if is_ole:
            try:
                import win32com.client
                import os as os_module
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(os_module.path.abspath(file_path))
                text = doc.Content.Text
                doc.Close(False)
                word.Quit()
                if text.strip():
                    logger.info(f"COM方式成功加载OLE文档: {file_path}")
                    return [Document(page_content=text, metadata={})]
            except Exception as e:
                logger.error(f"COM方式加载OLE文档失败: {str(e)}")
        
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            if text.strip():
                logger.info(f"备用方法成功加载DOCX: {file_path}")
                return [Document(page_content=text, metadata={})]
        except Exception as e:
            logger.warning(f"python-docx加载失败: {str(e)}")
        
        if not is_ole:
            try:
                import win32com.client
                import os as os_module
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(os_module.path.abspath(file_path))
                text = doc.Content.Text
                doc.Close(False)
                word.Quit()
                if text.strip():
                    logger.info(f"COM方式成功加载文档: {file_path}")
                    return [Document(page_content=text, metadata={})]
            except Exception as e:
                logger.error(f"所有方法都无法加载文档: {file_path}, 错误: {str(e)}")
        
        return []
    
    def load_pdf(self, file_path: str) -> List[Document]:
        """
        加载PDF格式文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            Document对象列表
        """
        try:
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            logger.info(f"成功加载PDF文档: {file_path}")
            return documents
        except Exception as e:
            logger.error(f"加载PDF文档失败: {file_path}, 错误: {str(e)}")
            return []
    
    def load_txt(self, file_path: str) -> List[Document]:
        """
        加载TXT格式文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            Document对象列表
        """
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
            logger.info(f"成功加载TXT文档: {file_path}")
            return documents
        except Exception as e:
            logger.error(f"加载TXT文档失败: {file_path}, 错误: {str(e)}")
            return []
    
    def load_doc(self, file_path: str) -> List[Document]:
        """
        加载DOC格式文档（旧版Word）
        
        Args:
            file_path: 文档路径
            
        Returns:
            Document对象列表
        """
        try:
            import win32com.client
            import os
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(os.path.abspath(file_path))
            text = doc.Content.Text
            doc.Close(False)
            word.Quit()
            
            documents = [Document(page_content=text, metadata={})]
            logger.info(f"成功加载DOC文档: {file_path}")
            return documents
        except Exception as e:
            logger.error(f"加载DOC文档失败: {file_path}, 错误: {str(e)}")
            return []
    
    def load_document(self, file_path: str) -> List[Document]:
        """
        根据文件扩展名自动选择加载器
        
        Args:
            file_path: 文档路径
            
        Returns:
            Document对象列表
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.docx':
            return self.load_docx(file_path)
        elif ext == '.pdf':
            return self.load_pdf(file_path)
        elif ext == '.txt':
            return self.load_txt(file_path)
        elif ext == '.doc':
            return self.load_doc(file_path)
        else:
            logger.warning(f"不支持的文件格式: {ext}")
            return []
    
    def data_cleaning(self, text: str) -> str:
        """
        数据清洗
        去除重复、无效内容，修正错别字，统一术语表述
        
        Args:
            text: 原始文本
            
        Returns:
            清洗后的文本
        """
        if not text:
            return ""
        
        text = re.sub(r'\s+', ' ', text)
        
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        text = text.strip()
        
        text = re.sub(r'第([一二三四五六七八九十百]+)条', lambda m: '第' + self._chinese_to_number(m.group(1)) + '条', text)
        
        return text
    
    def _chinese_to_number(self, chinese: str) -> str:
        """
        使用cn2an将中文数字转换为阿拉伯数字
        """
        try:
            return str(cn2an.cn2an(chinese, "normal"))
        except Exception:
            return chinese
    
    def extract_law_metadata(self, text: str, file_path: str) -> LegalMetadata:
        """
        从法条文本中提取元数据
        
        Args:
            text: 法条文本
            file_path: 文件路径
            
        Returns:
            LegalMetadata对象
        """
        file_name = os.path.basename(file_path)
        
        law_name_match = re.search(r'《(.+?)》', text)
        law_name = law_name_match.group(1) if law_name_match else os.path.splitext(file_name)[0]
        
        chapter_match = re.search(r'第([一二三四五六七八九十]+)[章节]', text)
        chapter = chapter_match.group(0) if chapter_match else "总则"
        
        article_number = ""
        article_patterns = [
            r'^第(\d+)条',
            r'\n第(\d+)条',
            r'。第(\d+)条',
        ]
        
        for pattern in article_patterns:
            article_match = re.search(pattern, text, re.MULTILINE)
            if article_match:
                num = article_match.group(1)
                if num and len(num) <= 3 and int(num) <= 500:
                    article_number = num
                    break
        
        return LegalMetadata(
            law_name=law_name,
            chapter=chapter,
            article_number=article_number,
            file_path=file_path,
            file_name=file_name
        )
    
    def extract_case_metadata(self, text: str, file_path: str) -> CaseMetadata:
        """
        从案例文本中提取元数据
        
        Args:
            text: 案例文本
            file_path: 文件路径
            
        Returns:
            CaseMetadata对象
        """
        file_name = os.path.basename(file_path)
        
        case_number_match = re.search(r'[（(](\d{4})[)）].*?第?(\d+号?)', text)
        case_number = case_number_match.group(0) if case_number_match else ""
        
        date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', text)
        judgment_date = date_match.group(0) if date_match else ""
        
        case_types = ['民事', '刑事', '行政', '经济', '知识产权', '劳动争议']
        case_type = "其他"
        for ct in case_types:
            if ct in text:
                case_type = ct
                break
        
        return CaseMetadata(
            case_number=case_number,
            judgment_date=judgment_date,
            case_type=case_type,
            file_path=file_path,
            file_name=file_name
        )
    
    def split_by_chapter(self, documents: List[Document], file_path: str, doc_type: str = "law") -> List[Document]:
        """
        按法条/章节拆分文档（父段）
        
        Args:
            documents: 原始文档列表
            file_path: 文件路径
            doc_type: 文档类型（law/case）
            
        Returns:
            按法条拆分后的文档列表
        """
        if not documents:
            return []
        
        full_text = "\n".join([doc.page_content for doc in documents])
        
        if doc_type == "law":
            split_docs = []
            current_chapter = "总则"
            
            chapter_match = re.search(r'第([一二三四五六七八九十]+)[章节]', full_text)
            if chapter_match:
                current_chapter = chapter_match.group(0)
            
            law_name_match = re.search(r'《(.+?)》', full_text)
            law_name = law_name_match.group(1) if law_name_match else os.path.splitext(os.path.basename(file_path))[0]
            
            article_pattern = r'第([一二三四五六七八九十百零]+)条'
            matches = list(re.finditer(article_pattern, full_text))
            
            if not matches:
                cleaned_text = self.data_cleaning(full_text)
                if len(cleaned_text) >= 30:
                    split_docs.append(Document(
                        page_content=cleaned_text,
                        metadata={
                            "law_name": law_name,
                            "chapter": current_chapter,
                            "article_number": "",
                            "case_number": "",
                            "judgment_date": "",
                            "case_type": "",
                            "file_path": file_path,
                            "file_name": os.path.basename(file_path),
                            "doc_type": doc_type
                        }
                    ))
            else:
                for i, match in enumerate(matches):
                    article_num = self._chinese_to_number(match.group(1))
                    start = match.start()
                    end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
                    article_text = full_text[start:end]
                    
                    cleaned_text = self.data_cleaning(article_text)
                    if len(cleaned_text) >= 30:
                        split_docs.append(Document(
                            page_content=cleaned_text,
                            metadata={
                                "law_name": law_name,
                                "chapter": current_chapter,
                                "article_number": article_num,
                                "case_number": "",
                                "judgment_date": "",
                                "case_type": "",
                                "file_path": file_path,
                                "file_name": os.path.basename(file_path),
                                "doc_type": doc_type
                            }
                        ))
        else:
            chapter_pattern = r'(【[^】]+】)'
            chapters = re.split(chapter_pattern, full_text)
            
            split_docs = []
            
            for i, chapter in enumerate(chapters):
                if not chapter.strip():
                    continue
                
                cleaned_text = self.data_cleaning(chapter)
                if len(cleaned_text) < 50:
                    continue
                
                metadata = self.extract_case_metadata(cleaned_text, file_path)
                
                split_docs.append(Document(
                    page_content=cleaned_text,
                    metadata={
                        "law_name": "",
                        "chapter": "",
                        "article_number": "",
                        "case_number": getattr(metadata, 'case_number', ''),
                        "judgment_date": getattr(metadata, 'judgment_date', ''),
                        "case_type": getattr(metadata, 'case_type', ''),
                        "file_path": file_path,
                        "file_name": os.path.basename(file_path),
                        "doc_type": doc_type
                    }
                ))
        
        return split_docs
    
    def split_to_chunks(self, documents: List[Document], doc_type: str = "law") -> List[Document]:
        """
        将父段拆分为子段
        
        Args:
            documents: 父段文档列表
            doc_type: 文档类型
            
        Returns:
            子段文档列表
        """
        chunks = []
        
        for doc in documents:
            sub_chunks = self.text_splitter.split_documents([doc])
            
            for i, chunk in enumerate(sub_chunks):
                chunk.metadata["chunk_index"] = i
                chunk.metadata["total_chunks"] = len(sub_chunks)
                chunks.append(chunk)
        
        return chunks
    
    def process_document(self, file_path: str, doc_type: str = "law") -> List[Document]:
        """
        处理单个文档的完整流程
        
        Args:
            file_path: 文档路径
            doc_type: 文档类型
            
        Returns:
            处理后的文档子段列表
        """
        documents = self.load_document(file_path)
        if not documents:
            return []
        
        chapter_docs = self.split_by_chapter(documents, file_path, doc_type)
        chunk_docs = self.split_to_chunks(chapter_docs, doc_type)
        
        logger.info(f"文档处理完成: {file_path}, 生成 {len(chunk_docs)} 个子段")
        return chunk_docs
    
    def process_directory(self, directory: str, doc_type: str = "law") -> List[Document]:
        """
        批量处理目录下的所有文档
        
        Args:
            directory: 目录路径
            doc_type: 文档类型
            
        Returns:
            所有文档的子段列表
        """
        all_chunks = []
        supported_extensions = ['.docx', '.pdf', '.txt']
        
        for root, dirs, files in os.walk(directory):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in supported_extensions:
                    file_path = os.path.join(root, file)
                    chunks = self.process_document(file_path, doc_type)
                    all_chunks.extend(chunks)
                    
                    if len(all_chunks) >= settings.BATCH_SIZE:
                        yield all_chunks
                        all_chunks = []
        
        if all_chunks:
            yield all_chunks


class EmbeddingGenerator:
    """
    嵌入向量生成器
    使用SentenceTransformer加载嵌入模型
    """
    
    def __init__(self, model_path: str = None):
        """
        初始化嵌入模型
        
        Args:
            model_path: 模型路径
        """
        self.model_path = model_path or settings.EMBEDDING_MODEL_PATH
        self.model = None
        
    def load_model(self):
        """
        加载嵌入模型
        """
        try:
            from sentence_transformers import SentenceTransformer
            
            logger.info(f"正在加载嵌入模型: {self.model_path}")
            self.model = SentenceTransformer(self.model_path)
            logger.info("嵌入模型加载完成")
            
        except Exception as e:
            logger.error(f"加载嵌入模型失败: {str(e)}")
            raise
    
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """
        批量生成嵌入向量
        
        Args:
            texts: 文本列表
            
        Returns:
            嵌入向量列表
        """
        if self.model is None:
            self.load_model()
        
        try:
            embeddings = self.model.encode(texts, show_progress_bar=False)
            return embeddings.tolist()
            
        except Exception as e:
            logger.error(f"生成嵌入向量失败: {str(e)}")
            raise
    
    def generate_single_embedding(self, text: str) -> List[float]:
        """
        生成单个文本的嵌入向量
        
        Args:
            text: 文本内容
            
        Returns:
            嵌入向量
        """
        return self.generate_embeddings([text])[0]


if __name__ == "__main__":
    processor = DocumentProcessor()
    embedding_generator = EmbeddingGenerator()
    
    test_dir = "./knowledge_base/laws"
    if os.path.exists(test_dir):
        for chunks in processor.process_directory(test_dir, "law"):
            texts = [chunk.page_content for chunk in chunks]
            embeddings = embedding_generator.generate_embeddings(texts)
            print(f"处理了 {len(chunks)} 个文档块")
